import os
from dotenv import load_dotenv
load_dotenv()
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

def extract_text_from_docx(path):
    try:
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"❌ Failed to parse DOCX: {path}, {e}")
        return ""

def load_or_build_vectorstore(debug=False):
    """Loads FAISS index from disk if available, else builds it from .docx and .txt files."""
    print("📂 Current working directory:", os.getcwd())
    print("🔍 Checking for FAISS index...")
    faiss_index_path = "faiss_index"
    embeddings = OpenAIEmbeddings()

    if os.path.exists(faiss_index_path):
        print("✅ FAISS index found. Loading from disk.")
        return FAISS.load_local(faiss_index_path, embeddings, allow_dangerous_deserialization=True)

    docs = []
    source_dir = "source_content"
    print("🧾 Files in source_content/:", os.listdir(source_dir) if os.path.isdir(source_dir) else "(Not found)")

    if os.path.isdir(source_dir):
        for file in os.listdir(source_dir):
            full_path = os.path.join(source_dir, file)
            if file.endswith(".docx"):
                print(f"📄 Found DOCX file: {file}")
                text = extract_text_from_docx(full_path)
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": file}))
            elif file.endswith(".txt"):
                print(f"📄 Found TXT file: {file}")
                try:
                    with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                        text = f.read()
                    print(f"✅ Read {len(text)} characters from {file}")
                    if text.strip():
                        docs.append(Document(page_content=text, metadata={"source": file}))
                except Exception as e:
                    print(f"❌ Failed to read {file}: {e}")

    if not docs:
        raise ValueError("🚨 No valid .docx or .txt files found in 'source_content/'")

    print(f"📚 Total documents parsed: {len(docs)}")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = splitter.split_documents(docs)
    print(f"🧩 Total chunks created: {len(chunks)}")

    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(faiss_index_path)
    print("✅ FAISS index built and saved.")
    return vectorstore
